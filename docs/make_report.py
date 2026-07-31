"""
Собирает итоговый отчёт (docs/report.docx) со всеми ключевыми инсайтами
проекта + графики. Разовый скрипт, не часть основного пайплайна.
"""
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path

OUT_DIR = Path(__file__).parent / "report_assets"
OUT_DIR.mkdir(exist_ok=True)

# --- график 1: сравнение моделей на UGC test ---
models = ["RF-DETR", "YOLO-seg", "Mask R-CNN", "SegFormer", "UNet-simple", "SAM zero-shot"]
ap50 = [0.277, 0.260, 0.203, 0.056, 0.050, 0.0002]
ap5095 = [0.175, 0.170, 0.131, 0.021, 0.022, 0.0002]

fig, ax = plt.subplots(figsize=(9, 5))
x = range(len(models))
ax.bar([i - 0.2 for i in x], ap50, width=0.4, label="segm AP@50")
ax.bar([i + 0.2 for i in x], ap5095, width=0.4, label="segm AP@50:95")
ax.set_xticks(list(x)); ax.set_xticklabels(models, rotation=20, ha="right")
ax.set_ylabel("AP"); ax.set_title("Сравнение моделей на UGC test (реальные фото Avito)")
ax.legend(); fig.tight_layout()
fig.savefig(OUT_DIR / "chart_models.png", dpi=130)
plt.close(fig)

# --- график 2: OCR эксперименты ---
ocr_labels = ["EasyOCR\nbaseline", "EasyOCR\n+CLAHE+mag", "Tesseract\ndigit-only", "PaddleOCR\n+CLAHE", "PaddleOCR\n(чистый)"]
ocr_recall = [13.4, 20.9, 19.4, 14.9, 38.8]
fig, ax = plt.subplots(figsize=(8, 5))
bars = ax.bar(ocr_labels, ocr_recall, color=["#888"]*4 + ["#2a7"])
ax.set_ylabel("Recall, %"); ax.set_title("OCR распознавания площади: recall по вариантам (67 размеченных комнат)")
for b, v in zip(bars, ocr_recall):
    ax.text(b.get_x()+b.get_width()/2, v+0.5, f"{v}%", ha="center")
fig.tight_layout()
fig.savefig(OUT_DIR / "chart_ocr.png", dpi=130)
plt.close(fig)

# --- график 3: domain gap (валидация vs UGC) для топ-3 моделей ---
fig, ax = plt.subplots(figsize=(7, 5))
labels = ["RF-DETR", "YOLO-seg", "Mask R-CNN"]
val_map = [0.834, 0.911, 0.858]  # segm mAP50 на валидации (синтетика)
ugc_map = [0.277, 0.260, 0.203]  # segm AP50 на UGC
x = range(len(labels))
ax.bar([i - 0.2 for i in x], val_map, width=0.4, label="Валидация (синтетика)")
ax.bar([i + 0.2 for i in x], ugc_map, width=0.4, label="UGC test (реальные фото)")
ax.set_xticks(list(x)); ax.set_xticklabels(labels)
ax.set_ylabel("segm mAP@50"); ax.set_title("Domain gap: валидация vs реальные фото")
ax.legend(); fig.tight_layout()
fig.savefig(OUT_DIR / "chart_domain_gap.png", dpi=130)
plt.close(fig)

# ===================== ДОКУМЕНТ =====================
doc = Document()
doc.add_heading("Сравнительное исследование instance segmentation для планов квартир", 0)
doc.add_paragraph("Задача: сегментация и классификация комнат на поэтажных планах квартир. "
                   "Обучение — на синтетических/чертёжных данных (ResPlan_v2 + CubiCasa5K), "
                   "тест — на реальных пользовательских фото с Avito/Cian (UGC).")

doc.add_heading("1. Ключевая находка: критический баг в разметке", level=1)
doc.add_paragraph(
    "На старте проекта была найдена и исправлена ошибка в вычислении bbox из полигонов масок: "
    "координаты считались из УПРОЩЁННОГО (approxPolyDP) контура вместо самого пиксельного массива "
    "маски. Это давало bbox МЕНЬШЕ реальной площади объекта у ~65% всех инстансов и у ~98% инстансов "
    "класса opening (двери/окна — тонкие вытянутые объекты, для которых упрощение контура особенно "
    "сильно искажает границы). Это единственная находка, которая повлияла на ВСЕ последующие модели "
    "и метрики — без фикса обучение шло на систематически неверных таргетах."
)

doc.add_heading("2. Валидация на обучающих данных (внутренний split, до UGC)", level=1)
doc.add_paragraph(
    "Метрики на 20% валидационной выборке из обучающих (чистых, схематичных) данных — "
    "показывают, насколько хорошо модель выучила саму задачу до переноса на реальные фото."
)
val_table = doc.add_table(rows=1, cols=5)
val_table.style = "Light Grid Accent 1"
vh = val_table.rows[0].cells
vh[0].text, vh[1].text, vh[2].text, vh[3].text, vh[4].text = (
    "Модель", "Эпоха", "box mAP50 / 50:95", "segm/mask mAP50 / 50:95", "своя метрика")
val_rows = [
    ("RF-DETR-Seg", "4", "0.900 / 0.826", "0.834 / 0.560", "—"),
    ("RF-DETR-Seg (новый прогон)", "9", "0.917 / 0.828", "0.844 / 0.578", "—"),
    ("YOLO-seg", "53", "0.942 / 0.877", "0.911 / 0.690", "—"),
    ("Mask R-CNN", "лучшая", "0.863 / 0.683", "0.858 / 0.675", "—"),
    ("SegFormer", "60 (финал)", "—", "—", "val_mIoU = 0.863"),
    ("UNet-simple", "24", "—", "—", "val_iou = 0.805"),
]
for row_vals in val_rows:
    row = val_table.add_row().cells
    for c, v in zip(row, row_vals):
        c.text = v
doc.add_paragraph(
    "Важное наблюдение: высокая метрика на валидации НЕ гарантирует качество на UGC — "
    "у SegFormer/UNet она даже выше, чем у YOLO/Mask R-CNN (mIoU/val_iou 0.80-0.86), но на "
    "реальных фото они проигрывают на порядок (см. раздел 3) — из-за отсутствия instance-head."
)

doc.add_heading("3. Итоговое сравнение моделей (UGC test, реальные фото)", level=1)
doc.add_picture(str(OUT_DIR / "chart_models.png"), width=Inches(6))
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Модель", "segm AP@50", "segm AP@50:95"
for m, a50, a95 in zip(models, ap50, ap5095):
    row = table.add_row().cells
    row[0].text, row[1].text, row[2].text = m, f"{a50:.4f}", f"{a95:.4f}"

doc.add_heading("4. Domain gap: чистая синтетика vs реальные фото", level=1)
doc.add_picture(str(OUT_DIR / "chart_domain_gap.png"), width=Inches(5.5))
doc.add_paragraph(
    "Все модели показывают провал в 3-5 раз между валидацией на обучающих (чистых, схематичных) "
    "данных и тестом на реальных UGC-фото (шум, компрессия, повороты, водяные знаки, разное качество "
    "съёмки). Это ожидаемый эффект domain gap, а не проблема конкретной архитектуры."
)
doc.add_paragraph(
    "Важный вывод: модели с нативным instance-head (RF-DETR, YOLO, Mask R-CNN) на порядок устойчивее "
    "моделей с чисто семантической сегментацией (SegFormer, UNet), даже когда у семантических моделей "
    "выше СВОЯ метрика (mIoU) на обучающих данных. У семантической сегментации нет объектной "
    "регуляризации — connected-components перевод в инстансы очень уязвим к шуму на реальных фото: "
    "рвёт один объект на несколько или сливает соседние объекты одного класса."
)

doc.add_heading("5. Что реально сделано против конкретных дефектов UGC", level=1)
doc.add_paragraph(
    "Важно честно разделить сделанное и обсуждённое:"
)
for bullet in [
    "Водяные знаки — реализован и протестирован эксперимент (data_prep/mask_watermark.py): "
    "закрашивание углов с логотипом Avito перед подачей в модель. Сравнение метрик "
    "YOLO до/после НЕ доведено до конца (переключились на другую задачу) — эффект пока не измерен.",
    "Повороты, JPEG-сжатие, разное качество съёмки — выявлены и задокументированы как причины "
    "domain gap (в т.ч. визуально при ручной разметке для OCR-эксперимента), но отдельных "
    "целевых фиксов под них не строили.",
    "Вместо точечных фиксов — сейчас тестируется набор аугментаций от коллеги "
    "(pipeline_fullaug_v2: 15 базовых + wall_styles/furniture/wall_morph_strong), обучение "
    "RF-DETR на них идёт в моменте написания отчёта — результат ещё не готов.",
]:
    doc.add_paragraph(bullet, style="List Bullet")

doc.add_heading("6. Zero-shot SAM практически не работает на этом домене", level=1)
doc.add_paragraph(
    "Grounded-SAM (GroundingDINO + SAM без дообучения) даёт AP≈0.0002 — на уровне шума. "
    "GroundingDINO обучен на натуральных фото и не умеет находить комнаты на схематичных чертежах — "
    "домен слишком далёк от его обучающих данных. Из 33 тестовых картинок находилось лишь 4 из 7 "
    "классов, с низкой уверенностью (0.26-0.44) и почти без пересечения по IoU с GT."
)

doc.add_heading("7. Furniture-aware реранкер: собранные данные и неудача генерализации", level=1)
doc.add_paragraph(
    "Собран мебельный датасет из двух источников: SFPI (189160 инстансов, готовый COCO-формат) и "
    "CubiCasa5K (28637 инстансов после исправления серьёзной ошибки в системе координат — SVG-полигоны "
    "мебели не совпадали с растровым изображением по системе координат; решение — калибровка через "
    "ECC-регистрацию изображений между отрендеренными SVG-стенами и реальным сканом)."
)
doc.add_paragraph(
    "YOLO11n-seg, обученный на этих данных, показывает отличные метрики на своей валидации "
    "(mask mAP50=0.99 на первой итерации), но НЕ ГЕНЕРАЛИЗУЕТСЯ на реальные UGC-фото — модель "
    "спамит класс 'sink' на любом мелком контрастном паттерне (100+ ложных срабатываний на "
    "картинку). Причина: оба источника данных — чистые CAD-иконки/векторная графика, визуально "
    "гораздо дальше от реальных фото, чем даже основной датасет комнат."
)

doc.add_heading("8. OCR-распознавание площади комнат", level=1)
doc.add_picture(str(OUT_DIR / "chart_ocr.png"), width=Inches(6))
doc.add_paragraph(
    "Идея: на планах печатают площадь комнаты (\"номер/площадь\" дробью по центру), это можно прочитать "
    "OCR и привязать к полигону комнаты геометрически (число внутри маски = площадь, а не размер стены). "
    "Протестированы EasyOCR, PaddleOCR и Tesseract (только цифры) с разными вариантами предобработки."
)
doc.add_paragraph(
    "PaddleOCR БЕЗ дополнительной обработки дал лучший результат — 38.8% recall на честной ручной "
    "разметке, почти втрое больше EasyOCR baseline (13.4%). Контрастность (CLAHE) помогает EasyOCR, "
    "но ВРЕДИТ PaddleOCR (у него уже есть своя настроенная предобработка). Печатный шрифт на части "
    "реальных фото физически нечитаем ни одним движком — это потолок исходного качества снимка, "
    "а не ограничение алгоритма."
)

doc.add_heading("9. Методологические решения", level=1)
for bullet in [
    "Единый train/valid split (80/20, seed=42) для всех моделей — честное сравнение.",
    "«Мягкая» (lenient) оценка: living/bedroom вместо room, bathroom+restroom объединены — "
    "не наказывать модель за семантически корректные, но формально иные названия.",
    "Confusion-матрица и раздельный учёт FP/FN по классам, а не только agregate mAP.",
    "GT-vs-предсказание визуализации для всех тестовых картинок, а не выборочно.",
]:
    doc.add_paragraph(bullet, style="List Bullet")

doc.add_heading("10. Открытые задачи", level=1)
for bullet in [
    "RF-DETR переобучается на новых аугментированных данных (wall_styles/furniture/wall_morph "
    "аугментации от коллеги) — ожидаем сокращение domain gap.",
    "FloorPlanCAD не подключён — не найден маппинг semantic-id → класс (список классов зашит в "
    "SVG-картинку как контуры, не как текст).",
    "SAM fine-tuning требует батчинга (уже добавлен в код) — предыдущая попытка не завершила "
    "даже одну эпоху за 4+ часа без него.",
    "Furniture-детектор нуждается в данных ближе к реальным фото или сильных аугментациях, "
    "иначе не может использоваться в реранкере.",
]:
    doc.add_paragraph(bullet, style="List Bullet")

import sys
out_name = sys.argv[1] if len(sys.argv) > 1 else "report.docx"
doc.save(str(Path(__file__).parent / out_name))
print("saved docs/report.docx")
